# src/parsers/document_parser.py

'''
Purpose of this file : This file will used to read and extract the content irrespective oof the format of the file

This file:

1) Accepts the file path
2) Detects the file type
3) Extracts the text from that file using pypdf, python-docx
4) Remove basic whitespace and newlines
5) Return the text and other meta data in below format

    Example output:\
    {
        "file_name": "resume.pdf", \
        "file_type": ".pdf", \
        "text": "Surabhi Nair ... Python ... BioVerify ...", \
        "char_count": 5421, \
        "word_count": 812 \
    }

'''

# --------------------------------------
# Importing Libraries
# --------------------------------------

from pypdf import PdfReader
from docx import Document
from pathlib import Path
import os
import re
from src.utils.text_cleaning import prepare_text_for_llm
## --------------------------------------
# Setup Configurations
# ---------------------------------------

# Set Paths
# BASE_DIR = Path.cwd().parent.parent
# DATA_DIR = BASE_DIR / "data"
# RAW_DATA = DATA_DIR / "raw"

# RESUME = RAW_DATA / "RESUME.DOCX"
# MOTIVATION =  RAW_DATA / "motivation.txt"
# PROFILE = RAW_DATA / "profile.md"
# ADDITIONAL = RAW_DATA / "additional.pdf"


# OUTPUT_DIR = DATA_DIR / "processed"
# OUTPUT_FILE = OUTPUT_DIR / "source.json"


# print(BASE_DIR)
# print(f"Data directory: {DATA_DIR} |  Data directory exists: {DATA_DIR.exists()}")
# print(f"Raw directory: {RAW_DATA} |  Raw directory exists: {RAW_DATA.exists()}")
# print(f"Output directory: {OUTPUT_DIR} |  Output directory exists: {OUTPUT_DIR.exists()}")
# print(f"Profile File: {PROFILE} |  Profile File exists: {PROFILE.exists()}")
# print(f"Resume File: {RESUME} |  Resume File exists: {RESUME.exists()}")
# print(f"Additional File: {ADDITIONAL} |  Additonal File exists: {ADDITIONAL.exists()}")
# print(f"Motivation File: {MOTIVATION} |  Motivation File exists: {MOTIVATION.exists()}")
# print(f"Output File: {OUTPUT_FILE} |  Output File exists: {OUTPUT_FILE.exists()}")

# --------------------------------------
# Helper Functions
# --------------------------------------

# Detecting the file type

def get_file_extension(path, name="Unknown") -> str:
    type_of_file = os.path.splitext(path)[1]
    # print(f"The type of the {name} is {type_of_file}")
    return type_of_file.lower()

# Convert into json format
def build_document_metadata(name, type_of_file, text, word_count, char_count) -> dict:
    return{
        "file_name": name,
        "file_type": type_of_file, 
        "text": text, 
        "char_count": char_count,
        "word_count": word_count
    }

# Clean the text
# def clean_extracted_text(text) -> str:
#     text = text.replace("\r\n", "\n").replace("\r", "\n")
#     text = re.sub(r"[ \t]+", " ", text)
#     text = re.sub(r"\n{3,}", "\n", text)
#     return text.strip()



# Docx handling function
def read_docx_file(path, name = "Unknown") -> str:
    text = Document(path)
    text_length = len(text.paragraphs)
    # print(f"{name} has been loaded | Para count is {text_length}")

    text_list = []

    for para in text.paragraphs:
        text = para.text.strip()
        text_list.append(text)

    text = "\n".join(text_list)

    return text


# PDF handling function
def read_pdf_file(path, name = "Unknown") -> str:
    text = PdfReader(path)
    text_length = len(text.pages)
    # print(f"{name} has been loaded | Pages count is {text_length}")

    text_list= []

    for page in text.pages:
        text = page.extract_text()
        if text:
            text_list.append(text.strip())

    text = "\n".join(text_list)

    # print(f"PDF Parsing completed | word count = {word_count}, char count = {char_count}")
    return text

    
# General File handling
def read_text_file(path, name = "Unknown", type_of_file=".txt") -> str:
    with open(path, "r", encoding="utf-8") as file:
        text = file.read().strip()

    # print(f"{name}{type_of_file} Parsing completed | word count = {word_count}, char count = {char_count}")

    return text

# Parse the final file
def parse_document(path, name="Unknown") -> dict:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError()
    type_of_file = get_file_extension(path, name)

    if type_of_file == ".pdf":
        text = read_pdf_file(path, name)
    elif type_of_file == ".docx":
        text = read_docx_file(path, name)
    elif type_of_file in [".txt", ".md"]:
        text = read_text_file(path, name, type_of_file)
    else:
        raise ValueError(f"Unsupported file type: ({type_of_file})")
    text = prepare_text_for_llm(text)
    if not text:
        raise ValueError(f"No text could be extracted from {path.name}")
    word_count = len(text.split())
    char_count = len(text)
    
    parsed_text = build_document_metadata(path.name, type_of_file, text, word_count, char_count)
    
    return parsed_text


# --------------------------------------
#Execute the functions
# resume = parse_file(RESUME, "RESUME")
# print(f"Output: {resume}")

# additional = parse_file(ADDITIONAL, "additional")
# print(f"Output: {additional}")

# profile = parse_file(PROFILE, 'profile')
# print(f"Output: {profile}")

# motivation =  parse_file(MOTIVATION, "Motivation")
# print(f"Output: {motivation}")


# --------------------------------------

# --------------------------------------
# Save the file
# --------------------------------------

# with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
#     json.dump(resume,f)
#     json.dump(additional,f)
#     json.dump(profile,f)
#     json.dump(motivation,f)